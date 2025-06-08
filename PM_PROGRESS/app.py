import re
import streamlit as st
import os
import json
import base64
import pypandoc
import tempfile
import shutil
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image
from auth.auth_manager import register_user, login_user
from utils.file_handler import create_project, load_user_projects
from lang_agent.proceed_l import generate_daily_log
import datetime

st.set_page_config(
    page_title="AI Project Work Tracker",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .main > div {
        padding-top: 2rem;
    }
    
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        text-align: center;
        color: white;
    }
    
    .card {
        background: white;
        padding: 2rem;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin-bottom: 1.5rem;
        border: 1px solid #e1e5e9;
    }
    
    .auth-container {
        max-width: 400px;
        margin: 0 auto;
        padding: 2rem;
        background: white;
        border-radius: 15px;
        box-shadow: 0 8px 25px rgba(0, 0, 0, 0.1);
        border: 1px solid #e1e5e9;
    }
    
    .stButton > button {
        width: 100%;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 0.75rem 1.5rem;
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4);
    }
    
    .css-1d391kg {
        background: linear-gradient(180deg, #f8f9fa 0%, #e9ecef 100%);
    }
    
    .stSuccess {
        background-color: #d4edda;
        border-color: #c3e6cb;
        color: #155724;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #28a745;
    }
    
    .stError {
        background-color: #f8d7da;
        border-color: #f5c6cb;
        color: #721c24;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #dc3545;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: #f8f9fa;
        border-radius: 8px;
        padding: 0.5rem 1rem;
        border: 1px solid #dee2e6;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
    }
    
    .stFileUploader > div {
        background-color: #f8f9fa;
        border: 2px dashed #dee2e6;
        border-radius: 8px;
        padding: 2rem;
        text-align: center;
    }
    
    .stTextArea > div > div > textarea {
        border-radius: 8px;
        border: 2px solid #e1e5e9;
    }
    
    .stSelectbox > div > div {
        border-radius: 8px;
        border: 2px solid #e1e5e9;
    }
    
    .progress-indicator {
        display: flex;
        justify-content: center;
        align-items: center;
        padding: 1rem;
        background: #f8f9fa;
        border-radius: 8px;
        margin: 1rem 0;
    }
    
    .welcome-message {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 1rem;
        font-weight: 600;
    }
    
    .nav-item {
        padding: 0.5rem;
        margin: 0.25rem 0;
        border-radius: 8px;
        transition: all 0.3s ease;
    }
    
    .image-preview {
        border-radius: 8px;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1);
        margin: 1rem 0;
    }
    
    .stDownloadButton > button {
        background: linear-gradient(90deg, #28a745 0%, #20c997 100%);
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 6px;
        font-weight: 500;
        margin: 0.25rem;
    }
</style>
""", unsafe_allow_html=True)

USER_LOG_PATH = "data/log"

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "auth_attempted" not in st.session_state:
    st.session_state.auth_attempted = False

def convert_md_to_docx(markdown_content, user_dir, output_filename):
    try:
        pypandoc.download_pandoc()
    except:
        pass
    
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as temp_md:
            temp_md.write(markdown_content)
            temp_md_path = temp_md.name
        
        temp_docx_path = tempfile.mktemp(suffix='.docx')
        
        pypandoc.convert_file(temp_md_path, 'docx', outputfile=temp_docx_path)
        
        final_output_path = os.path.join(user_dir, output_filename)
        shutil.move(temp_docx_path, final_output_path)
        
        os.unlink(temp_md_path)
        
        return final_output_path
    except:
        return convert_md_to_docx_manual(markdown_content, user_dir, output_filename)

def convert_md_to_docx_manual(markdown_content, user_dir, output_filename):
    doc = Document()
    
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('*') and line.endswith('*'):
            p = doc.add_paragraph()
            run = p.add_run(line[1:-1])
            run.italic = True
        elif line.startswith('![') and '](' in line and line.endswith(')'):
            img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                abs_img_path = os.path.join(user_dir, img_path)
                if os.path.exists(abs_img_path):
                    try:
                        paragraph = doc.add_paragraph()
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        doc.add_picture(abs_img_path, width=Inches(4))
                        if alt_text:
                            caption = doc.add_paragraph(alt_text)
                            caption.alignment = 1
                    except:
                        doc.add_paragraph(f"[Image: {alt_text or os.path.basename(img_path)}]")
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            doc.add_paragraph(line[3:], style='List Number')
        else:
            if line:
                doc.add_paragraph(line)
    
    output_path = os.path.join(user_dir, output_filename)
    doc.save(output_path)
    return output_path

def get_download_button(file_path, filename, label):
    if os.path.exists(file_path):
        with open(file_path, "rb") as file:
            file_data = file.read()
        st.download_button(
            label=label,
            data=file_data,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        return True
    return False

st.markdown("""
<div class="main-header">
    <h1>🧠 AI Project Work Tracker</h1>
    <p>Track your daily progress with AI-powered insights</p>
</div>
""", unsafe_allow_html=True)

def handle_login(username, password):
    if login_user(username, password):
        st.session_state.logged_in = True
        st.session_state.username = username
        st.session_state.auth_attempted = True
        st.rerun()
    else:
        st.session_state.auth_attempted = True
        return False
    return True

def handle_register(username, email, password):
    if register_user(username, email, password):
        st.session_state.auth_attempted = True
        return True
    else:
        st.session_state.auth_attempted = True
        return False

if not st.session_state.logged_in:
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown('<div class="auth-container">', unsafe_allow_html=True)
        
        auth_mode = st.radio(
            "🔐 Authentication Mode", 
            ["Login", "Register"],
            horizontal=True,
            help="Choose whether to login or register a new account"
        )
        
        st.markdown("---")
        
        username = st.text_input(
            "👤 Username", 
            placeholder="Enter your username",
            help="Your unique username"
        )
        password = st.text_input(
            "🔒 Password", 
            type="password",
            placeholder="Enter your password",
            help="Your secure password"
        )
        
        email = ""
        if auth_mode == "Register":
            email = st.text_input(
                "📧 Email", 
                placeholder="Enter your email address",
                help="Your email address for account recovery"
            )
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        if st.button(f"🚀 {auth_mode}", key="auth_button"):
            if auth_mode == "Login":
                if username and password:
                    with st.spinner("Logging in..."):
                        if not handle_login(username, password):
                            st.error("❌ Invalid credentials. Please try again.")
                else:
                    st.warning("⚠️ Please fill in all fields")
            else:  
                if username and password and email:
                    with st.spinner("Creating account..."):
                        if handle_register(username, email, password):
                            st.success("✅ Registration successful! Please login.")
                        else:
                            st.error("❌ Username already exists. Please choose another.")
                else:
                    st.warning("⚠️ Please fill in all fields")
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("---")
        if auth_mode == "Login":
            st.info("💡 Don't have an account? Select 'Register' above to create one.")
        else:
            st.info("💡 Already have an account? Select 'Login' above to sign in.")

else:
    with st.sidebar:
        st.markdown(f"""
        <div class="welcome-message">
            👋 Welcome back, <strong>{st.session_state.username}</strong>!
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("### 🧭 Navigation")
        menu_options = [
            "➕ Add Today's Update",
            "📜 View Logs", 
            "📁 Add Project"
        ]
        
        menu = st.radio(
            "Choose an option:",
            menu_options,
            help="Select what you'd like to do"
        )
        
        st.markdown("---")
        
        if st.button("🚪 Logout", help="Sign out of your account"):
            st.session_state.logged_in = False
            st.session_state.username = ""
            st.session_state.auth_attempted = False
            st.rerun()
    
    if menu == "➕ Add Today's Update":
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.subheader("📝 Add Today's Project Log")
        st.markdown("Record your daily progress and achievements")
        st.markdown('</div>', unsafe_allow_html=True)

        projects = load_user_projects(st.session_state.username)
        if not projects:
            st.warning("⚠️ You have no projects yet. Please add one first using the 'Add Project' section.")
        else:
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.markdown('<div class="card">', unsafe_allow_html=True)
                selected_project = st.selectbox(
                    "🎯 Select Project", 
                    [p['project_name'] for p in projects],
                    help="Choose the project you worked on today"
                )
                selected_project_desc = next((p['description'] for p in projects if p['project_name'] == selected_project), "")
                
                if selected_project_desc:
                    st.info(f"📋 **Project Description:** {selected_project_desc}")
                
                user_input = st.text_area(
                    "✍️ What did you accomplish today?",
                    height=150,
                    placeholder="Describe your work, challenges faced, solutions found, and next steps...",
                    help="Be detailed about your progress, it helps generate better logs"
                )
                st.markdown('</div>', unsafe_allow_html=True)
            
            with col2:
                st.markdown('<div class="card">', unsafe_allow_html=True)
                st.markdown("📸 **Upload Screenshots**")
                uploaded_images = st.file_uploader(
                    "Upload any screenshots or images",
                    accept_multiple_files=True,
                    type=['png', 'jpg', 'jpeg', 'gif'],
                    help="Upload images to document your progress visually"
                )
                
                if uploaded_images:
                    st.success(f"✅ {len(uploaded_images)} image(s) selected")
                st.markdown('</div>', unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                if st.button("🤖 Generate & Save Log", key="generate_log", help="AI will process your input and create a detailed log"):
                    if user_input.strip():
                        with st.spinner("🔄 Generating your daily log..."):
                            user_dir = os.path.join(USER_LOG_PATH, st.session_state.username)
                            os.makedirs(user_dir, exist_ok=True)
                            md_file = os.path.join(user_dir, f"{st.session_state.username}_logs.md")
                            json_file = os.path.join(user_dir, f"{st.session_state.username}_logs.json")

                            img_paths = []
                            if uploaded_images:
                                img_dir = os.path.join(user_dir, "assets")
                                os.makedirs(img_dir, exist_ok=True)
                                for img in uploaded_images:
                                    img_path = os.path.join(img_dir, img.name)
                                    with open(img_path, "wb") as f:
                                        f.write(img.read())
                                    img_paths.append(img_path)

                            img_md_lines = [f"![{os.path.basename(p)}](assets/{os.path.basename(p)})" for p in img_paths]
                            img_md_block = "\n".join(img_md_lines)

                            current_date = datetime.datetime.now().strftime("%Y-%m-%d")

                            try:
                                new_log = generate_daily_log(
                                    selected_project, selected_project_desc, user_input, 
                                    st.session_state.username, current_date, img_paths,
                                )
                            except:
                                new_log = f"### 📊 {selected_project} - {current_date}\n\n**Today's Progress:**\n{user_input}"

                            with open(md_file, "a") as f:
                                f.write(f"\n\n{new_log}\n\n{img_md_block}")

                            if os.path.exists(json_file):
                                with open(json_file, "r") as f:
                                    logs_dict = json.load(f)
                            else:
                                logs_dict = {}
                            
                            logs_dict[current_date] = f"{new_log}\n\n{img_md_block}"
                            
                            with open(json_file, "w") as f:
                                json.dump(logs_dict, f, indent=2)

                        st.success("🎉 Log saved successfully!")
                        st.markdown("### 📋 Generated Log Preview")
                        
                        st.markdown('<div class="card">', unsafe_allow_html=True)
                        log_lines = new_log.split('\n')
                        for line in log_lines:
                            img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
                            if img_match:
                                alt_text = img_match.group(1)
                                img_path = img_match.group(2)
                                abs_img_path = os.path.join(USER_LOG_PATH, st.session_state.username, img_path)
                                if os.path.exists(abs_img_path):
                                    try:
                                        with open(abs_img_path, "rb") as image_file:
                                            base64_string = base64.b64encode(image_file.read()).decode('utf-8')
                                        mime_type = "image/png" if img_path.lower().endswith('.png') else "image/jpeg"
                                        data_url = f"data:{mime_type};base64,{base64_string}"
                                        st.image(data_url, caption=alt_text if alt_text else os.path.basename(img_path), use_column_width=True)
                                    except Exception as e:
                                        st.warning(f"⚠️ Failed to load image {os.path.basename(img_path)}: {str(e)}")
                            else:
                                if line.strip():
                                    st.markdown(line.strip())
                        st.markdown('</div>', unsafe_allow_html=True)
                    else:
                        st.warning("⚠️ Please describe what you did today before generating the log.")

    elif menu == "📜 View Logs":
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.subheader("📚 View Your Project Logs")
        st.markdown("Browse through your daily progress and achievements")
        st.markdown('</div>', unsafe_allow_html=True)

        tabs = st.tabs(["📄 Complete Log History", "📅 Specific Date"])
        user_dir = os.path.join(USER_LOG_PATH, st.session_state.username)
        md_file = os.path.join(user_dir, f"{st.session_state.username}_logs.md")
        json_file = os.path.join(user_dir, f"{st.session_state.username}_logs.json")

        with tabs[0]:
            st.markdown("### 📖 Your Complete Work History")
            
            col1, col2 = st.columns([3, 1])
            with col2:
                if os.path.exists(md_file):
                    if st.button("📥 Download Complete Report", key="download_all"):
                        with st.spinner("Generating Word document..."):
                            try:
                                with open(md_file, "r") as f:
                                    full_content = f.read()
                                
                                docx_filename = f"{st.session_state.username}_complete_logs.docx"
                                docx_path = convert_md_to_docx(full_content, user_dir, docx_filename)
                                
                                get_download_button(docx_path, docx_filename, "📄 Download Word Report")
                                st.success("Document ready for download!")
                            except Exception as e:
                                st.error(f"Error generating document: {str(e)}")
            
            with col1:
                if os.path.exists(md_file):
                    with open(md_file, "r") as f:
                        logs = f.readlines()
                    
                    st.markdown('<div class="card">', unsafe_allow_html=True)
                    for line in logs:
                        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
                        if img_match:
                            alt_text = img_match.group(1)
                            img_path = img_match.group(2)
                            abs_img_path = os.path.join(USER_LOG_PATH, st.session_state.username, img_path)
                            if os.path.exists(abs_img_path):
                                try:
                                    with open(abs_img_path, "rb") as image_file:
                                        base64_string = base64.b64encode(image_file.read()).decode('utf-8')
                                    mime_type = "image/png" if img_path.lower().endswith('.png') else "image/jpeg"
                                    data_url = f"data:{mime_type};base64,{base64_string}"
                                    st.image(data_url, caption=alt_text if alt_text else os.path.basename(img_path), use_column_width=True)
                                except Exception as e:
                                    st.warning(f"⚠️ Failed to load image {os.path.basename(img_path)}: {str(e)}")
                        else:
                            if line.strip():
                                st.markdown(line.strip())
                    st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.info("📝 No logs found yet. Start by adding your first daily update!")

        with tabs[1]:
            st.markdown("### 📅 View Specific Date")
            if os.path.exists(json_file):
                with open(json_file, "r") as f:
                    logs_dict = json.load(f)
                
                col1, col2 = st.columns([1, 2])
                with col1:
                    available_dates = list(logs_dict.keys())
                    st.info(f"📊 Available dates: {len(available_dates)}")
                    selected_date = st.date_input(
                        "🗓️ Select a date to view log:",
                        help="Choose a date to view your work log for that day"
                    )
                    
                    selected_date_str = selected_date.strftime("%Y-%m-%d")
                    if selected_date_str in logs_dict:
                        if st.button("📥 Download Date Report", key="download_date"):
                            with st.spinner("Generating Word document..."):
                                try:
                                    date_content = logs_dict[selected_date_str]
                                    docx_filename = f"{st.session_state.username}_{selected_date_str}_log.docx"
                                    docx_path = convert_md_to_docx(date_content, user_dir, docx_filename)
                                    
                                    get_download_button(docx_path, docx_filename, f"📄 Download {selected_date_str} Report")
                                    st.success("Document ready for download!")
                                except Exception as e:
                                    st.error(f"Error generating document: {str(e)}")
                
                with col2:
                    if selected_date_str in logs_dict:
                        st.markdown(f"### 📋 Log for {selected_date_str}")
                        st.markdown('<div class="card">', unsafe_allow_html=True)
                        log_lines = logs_dict[selected_date_str].split('\n')
                        for line in log_lines:
                            img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
                            if img_match:
                                alt_text = img_match.group(1)
                                img_path = img_match.group(2)
                                abs_img_path = os.path.join(USER_LOG_PATH, st.session_state.username, img_path)
                                if os.path.exists(abs_img_path):
                                    try:
                                        with open(abs_img_path, "rb") as image_file:
                                            base64_string = base64.b64encode(image_file.read()).decode('utf-8')
                                        mime_type = "image/png" if img_path.lower().endswith('.png') else "image/jpeg"
                                        data_url = f"data:{mime_type};base64,{base64_string}"
                                        st.image(data_url, caption=alt_text if alt_text else os.path.basename(img_path), use_column_width=True)
                                    except Exception as e:
                                        st.warning(f"⚠️ Failed to load image {os.path.basename(img_path)}: {str(e)}")
                            else:
                                if line.strip():
                                    st.markdown(line.strip())
                        st.markdown('</div>', unsafe_allow_html=True)
                    else:
                        st.warning("⚠️ No log found for the selected date.")
            else:
                st.info("📝 No logs found yet. Start by adding your first daily update!")

    elif menu == "📁 Add Project":
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.subheader("🚀 Create a New Project")
        st.markdown("Set up a new project to start tracking your daily progress")
        st.markdown('</div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            project_name = st.text_input(
                "📋 Project Name",
                placeholder="Enter a clear, descriptive project name",
                help="Choose a name that clearly identifies your project"
            )
            project_desc = st.text_area(
                "📝 Project Description",
                height=150,
                placeholder="Describe the project goals, scope, and key objectives...",
                help="Provide a detailed description to help generate better daily logs"
            )
            st.markdown('</div>', unsafe_allow_html=True)
        
        with col2:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown("💡 **Tips for Success:**")
            st.markdown("""
            - Use clear, specific project names
            - Include project goals in description
            - Mention key technologies/tools
            - Define success criteria
            - Note important deadlines
            """)
            st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🎯 Create Project", key="create_project", help="Add this project to your workspace"):
                if project_name.strip() and project_desc.strip():
                    with st.spinner("Creating project..."):
                        if create_project(st.session_state.username, project_name, project_desc):
                            st.success("🎉 Project created successfully! You can now start logging daily updates.")
                        else:
                            st.error("❌ A project with this name already exists. Please choose a different name.")
                else:
                    st.warning("⚠️ Please fill in both project name and description.")

        st.markdown("---")
        st.markdown("### 📂 Your Existing Projects")
        projects = load_user_projects(st.session_state.username)
        if projects:
            for i, project in enumerate(projects):
                with st.expander(f"📋 {project['project_name']}", expanded=False):
                    st.markdown(f"**Description:** {project['description']}")
        else:
            st.info("📝 No projects yet. Create your first project above!")